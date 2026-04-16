# file_io_server.py
import os
import re
import base64
import json
import time
import requests
from pathlib import Path
from typing import Optional, Union
from fastmcp import FastMCP

# Optional imports for extended formats
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook, Workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

# Configuration: base directory for file operations
BASE_DIR = os.environ.get("FILE_IO_BASE_DIR", "./workspace")
os.makedirs(BASE_DIR, exist_ok=True)

mcp = FastMCP("File I/O Server")

# -------------------------------------------------------------------------
# Helper functions
# -------------------------------------------------------------------------
def detect_extension(content: str) -> str:
    """Detect file extension based on content heuristics."""
    content_lower = content.lower()
    # Python
    if re.search(r'^\s*import\s+\w+|^\s*from\s+\w+\s+import|^\s*def\s+\w+\s*\(|^\s*class\s+\w+', content, re.MULTILINE):
        return ".py"
    # JavaScript
    if re.search(r'^\s*(const|let|var|function|=>|console\.log|document\.|window\.)', content, re.MULTILINE):
        return ".js"
    # HTML
    if re.search(r'^\s*<!DOCTYPE\s+html|<html|<body|<div', content, re.MULTILINE):
        return ".html"
    # CSS
    if re.search(r'^\s*[a-zA-Z-]+\s*{[^}]*}', content, re.MULTILINE):
        return ".css"
    # JSON
    if content.strip().startswith('{') and content.strip().endswith('}'):
        try:
            json.loads(content)
            return ".json"
        except:
            pass
    # Markdown
    if re.search(r'^#{1,6}\s+\w+|^\s*[-*+]\s+\w+|\[.*\]\(.*\)', content, re.MULTILINE):
        return ".md"
    return ".txt"

def read_text_file(path: Path) -> str:
    """Read a text file with utf-8 encoding."""
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def write_text_file(path: Path, content: str) -> str:
    """Write a text file with utf-8 encoding."""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    return f"Successfully wrote text file to {path}"

def read_docx(path: Path) -> str:
    """Extract text from a .docx file."""
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Run: pip install python-docx"
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def write_docx(path: Path, content: str) -> str:
    """Create a .docx file with the given text content."""
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Run: pip install python-docx"
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(path)
    return f"Successfully wrote .docx file to {path}"

def read_xlsx(path: Path) -> str:
    """Read a .xlsx file and return as text (sheet by sheet)."""
    if not XLSX_AVAILABLE:
        return "Error: openpyxl not installed. Run: pip install openpyxl"
    wb = load_workbook(path, data_only=True)
    output = []
    for sheet in wb.worksheets:
        output.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            output.append("\t".join(str(cell) if cell is not None else "" for cell in row))
    return "\n".join(output)

def write_xlsx(path: Path, content: str) -> str:
    """Create a .xlsx file from tabular text (rows separated by newline, cells by tabs)."""
    if not XLSX_AVAILABLE:
        return "Error: openpyxl not installed. Run: pip install openpyxl"
    wb = Workbook()
    ws = wb.active
    rows = content.strip().split('\n')
    for i, row in enumerate(rows, 1):
        cells = row.split('\t')
        for j, cell in enumerate(cells, 1):
            ws.cell(row=i, column=j, value=cell)
    wb.save(path)
    return f"Successfully wrote .xlsx file to {path}"

def read_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    if not PDF_AVAILABLE:
        return "Error: pypdf or PyPDF2 not installed. Run: pip install pypdf"
    reader = PdfReader(path)
    text = []
    for page in reader.pages:
        text.append(page.extract_text())
    return "\n".join(text)

def write_pdf(path: Path, content: str) -> str:
    """Create a simple PDF from text using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return "Error: reportlab not installed. Run: pip install reportlab"
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 40
    for line in content.split('\n'):
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line[:100])
        y -= 14
    c.save()
    return f"Successfully wrote PDF file to {path}"

# -------------------------------------------------------------------------
# Tools
# -------------------------------------------------------------------------
@mcp.tool()
def read_file(file_path: str) -> str:
    """
    Read the content of a file. Supports .txt, .py, .js, .json, .md, .html, .css,
    .docx, .xlsx, .pdf. Returns the extracted text.
    """
    path = Path(BASE_DIR) / file_path
    if not path.exists():
        return f"Error: File not found: {path}"
    
    ext = path.suffix.lower()
    try:
        if ext == '.docx':
            return read_docx(path)
        elif ext == '.xlsx':
            return read_xlsx(path)
        elif ext == '.pdf':
            return read_pdf(path)
        else:
            # Treat as text
            return read_text_file(path)
    except Exception as e:
        return f"Error reading file: {e}"

@mcp.tool()
def write_file(content: str, filename: Optional[str] = None, extension: Optional[str] = None) -> str:
    """
    Write content to a file. Automatically detects extension from content if not provided.
    Supports .txt, .py, .js, .json, .md, .html, .css, .docx, .xlsx, .pdf.
    For binary formats (.docx, .xlsx, .pdf), the content should be plain text; the server
    will generate the appropriate binary file.
    """
    # Determine extension
    if extension:
        ext = extension if extension.startswith('.') else f'.{extension}'
    else:
        ext = detect_extension(content)
    
    # Generate filename if not provided
    if not filename:
        filename = f"file_{int(time.time())}"
    # Remove any existing extension from filename
    if '.' in filename:
        filename = filename.split('.')[0]
    
    full_path = Path(BASE_DIR) / f"{filename}{ext}"
    
    # Write based on extension
    try:
        if ext == '.docx':
            return write_docx(full_path, content)
        elif ext == '.xlsx':
            return write_xlsx(full_path, content)
        elif ext == '.pdf':
            return write_pdf(full_path, content)
        else:
            return write_text_file(full_path, content)
    except Exception as e:
        return f"Error writing file: {e}"


@mcp.tool()
def perform_websearch(query: str) -> str:
    """
    Performs a web search for a query
    Args:
        query: the query to web search.
    Returns:
        A formatted string with search results
    """
    
    try:
        params = {
            'access_key': YOUR_API_KEY,
            'query': query
        }

        api_result = requests.get('https://api.serpstack.com/search', params=params)
        api_result.raise_for_status()  # Raise exception for bad status codes
        
        api_response = api_result.json()

        # Check if search was successful
        if 'error' in api_response:
            return f"Error: {api_response['error']['info']}"

        # Format results
        results_text = f"Total results: {api_response.get('search_information', {}).get('total_results', 'N/A')}\n\n"
        
        if 'organic_results' in api_response:
            for number, result in enumerate(api_response['organic_results'], start=1):
                title = result.get('title', 'No title')
                link = result.get('url', 'No URL')
                snippet = result.get('snippet', 'No snippet')
                results_text += f"{number}. {title}\n"
                results_text += f"   URL: {link}\n"
                results_text += f"   {snippet}\n\n"
        else:
            results_text += "No organic results found."
        
        return results_text
    
    except requests.exceptions.RequestException as e:
        return f"Request error: {str(e)}"
    except Exception as e:
        return f"Error: {str(e)}"
        
@mcp.tool()
def greeting(name: str) -> str:
    """Send a greeting"""
    try:
        if not name:
            raise ValueError("Name cannot be empty")
        return f"Hi, how are you, {name}"
    except Exception as e:
        #logger.error(f"Tool execution failed: {e}")
        raise e
    
@mcp.resource("files://list")
def list_files() -> str:
    """List all files in the workspace directory."""
    try:
        files = [f for f in os.listdir(BASE_DIR) if os.path.isfile(os.path.join(BASE_DIR, f))]
        if not files:
            return "No files found."
        return "\n".join(files)
    except Exception as e:
        return f"Error listing files: {e}"
